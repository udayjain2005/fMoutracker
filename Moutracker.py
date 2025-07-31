from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
import openpyxl
from io import BytesIO
import os
from werkzeug.utils import secure_filename
from docx import Document
import json
import sys
from jinja2 import pass_context

#uday jain code
app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///moutracker.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'your_secret_key'
db = SQLAlchemy(app)

UPLOAD_FOLDER = 'templates/word_templates'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
generated_docs_folder = 'generated_docs'
os.makedirs(generated_docs_folder, exist_ok=True)

# MasterValue table for all dropdowns
class MasterValue(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False)
    category = db.Column(db.String(50), nullable=False)  # e.g., 'type', 'sector', etc.
    __table_args__ = (db.UniqueConstraint('name', 'category', name='_name_category_uc'),)

# Models
class MoU(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(120), nullable=False)  # New column for Title
    start_date = db.Column(db.Date, nullable=False)
    end_date = db.Column(db.Date, nullable=False)
    company_name = db.Column(db.String(120), nullable=False)
    type_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    sector_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    investment = db.Column(db.Float, nullable=False)
    country_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    state_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    district_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    stage_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    land_status_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    clearance_status_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    commissioning_status_id = db.Column(db.Integer, db.ForeignKey('master_value.id'))
    remarks = db.Column(db.Text)
    # Relationships
    type = db.relationship('MasterValue', foreign_keys=[type_id])
    sector = db.relationship('MasterValue', foreign_keys=[sector_id])
    country = db.relationship('MasterValue', foreign_keys=[country_id])
    state = db.relationship('MasterValue', foreign_keys=[state_id])
    district = db.relationship('MasterValue', foreign_keys=[district_id])
    stage = db.relationship('MasterValue', foreign_keys=[stage_id])
    land_status = db.relationship('MasterValue', foreign_keys=[land_status_id])
    clearance_status = db.relationship('MasterValue', foreign_keys=[clearance_status_id])
    commissioning_status = db.relationship('MasterValue', foreign_keys=[commissioning_status_id])

class AuditLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    action_type = db.Column(db.String(20))
    mou_id = db.Column(db.Integer)
    description = db.Column(db.Text)

# Jinja2 file exists filter
@pass_context
def file_exists(context, path):
    return os.path.exists(path)

app.jinja_env.filters['file_exists'] = file_exists

# Routes
@app.route('/')
def index():
    return redirect(url_for('list_mous'))

@app.route('/mous')
def list_mous():
    # Filtering
    query = MoU.query
    company = request.args.get('company', '').strip()
    country_id = request.args.get('country', '').strip()
    type_id = request.args.get('type', '').strip()
    sector_id = request.args.get('sector', '').strip()
    if company:
        query = query.filter(MoU.company_name.ilike(f'%{company}%'))
    if country_id:
        query = query.filter(MoU.country_id == int(country_id))
    if type_id:
        query = query.filter(MoU.type_id == int(type_id))
    if sector_id:
        query = query.filter(MoU.sector_id == int(sector_id))
    # Sorting
    sort = request.args.get('sort', 'id')
    direction = request.args.get('direction', 'asc')
    if hasattr(MoU, sort):
        sort_col = getattr(MoU, sort)
        if direction == 'desc':
            sort_col = sort_col.desc()
        query = query.order_by(sort_col)
    # Pagination
    page = request.args.get('page', 1, type=int)
    per_page = 10
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    mous = pagination.items
    # For filter dropdowns (use MasterValue)
    countries = MasterValue.query.filter_by(category='country').order_by(MasterValue.name).all()
    types = MasterValue.query.filter_by(category='type').order_by(MasterValue.name).all()
    sectors = MasterValue.query.filter_by(category='sector').order_by(MasterValue.name).all()
    # Format dates for display
    for mou in mous:
        mou.start_date_fmt = mou.start_date.strftime('%d-%m-%Y') if mou.start_date else ''
        mou.end_date_fmt = mou.end_date.strftime('%d-%m-%Y') if mou.end_date else ''
    return render_template('list.html', mous=mous, pagination=pagination, countries=countries, types=types, sectors=sectors, request=request)

@app.route('/mou/new', methods=['GET', 'POST'])
def create_mou():
    types = MasterValue.query.filter_by(category='type').order_by(MasterValue.name).all()
    sectors = MasterValue.query.filter_by(category='sector').order_by(MasterValue.name).all()
    countries = MasterValue.query.filter_by(category='country').order_by(MasterValue.name).all()
    states = MasterValue.query.filter_by(category='state').order_by(MasterValue.name).all()
    districts = MasterValue.query.filter_by(category='district').order_by(MasterValue.name).all()
    stages = MasterValue.query.filter_by(category='stage').order_by(MasterValue.name).all()
    land_statuses = MasterValue.query.filter_by(category='land_status').order_by(MasterValue.name).all()
    clearance_statuses = MasterValue.query.filter_by(category='clearance_status').order_by(MasterValue.name).all()
    commissioning_statuses = MasterValue.query.filter_by(category='commissioning_status').order_by(MasterValue.name).all()
    if request.method == 'POST':
        # Validate and create MoU
        title = request.form.get('Title', '').strip()
        company_name = request.form.get('company_name', '').strip()
        type_id = request.form.get('type_id', type=int)
        sector_id = request.form.get('sector_id', type=int)
        investment = request.form.get('investment', type=float)
        country_id = request.form.get('country_id', type=int)
        state_id = request.form.get('state_id', type=int)
        district_id = request.form.get('district_id', type=int)
        stage_id = request.form.get('stage_id', type=int)
        land_status_id = request.form.get('land_status_id', type=int)
        clearance_status_id = request.form.get('clearance_status_id', type=int)
        commissioning_status_id = request.form.get('commissioning_status_id', type=int)
        start_date = datetime.strptime(request.form.get('start_date', ''), '%d-%m-%Y').date()
        end_date = datetime.strptime(request.form.get('end_date', ''), '%d-%m-%Y').date()
        remarks = request.form.get('remarks', '').strip()
        mou = MoU(
            title=title,
            company_name=company_name,
            type_id=type_id,
            sector_id=sector_id,
            investment=investment,
            country_id=country_id,
            state_id=state_id,
            district_id=district_id,
            stage_id=stage_id,
            land_status_id=land_status_id,
            clearance_status_id=clearance_status_id,
            commissioning_status_id=commissioning_status_id,
            start_date=start_date,
            end_date=end_date,
            remarks=remarks
        )
        db.session.add(mou)
        db.session.commit()
        # Audit log
        db.session.add(AuditLog(
            action_type='CREATE',
            mou_id=mou.id,
            description=f"Created MoU: {mou.title} for {mou.company_name}"
        ))
        db.session.commit()
        flash('MoU created successfully!', 'success')
        return redirect(url_for('list_mous'))
    return render_template('form.html', types=types, sectors=sectors, countries=countries, states=states, districts=districts, stages=stages, land_statuses=land_statuses, clearance_statuses=clearance_statuses, commissioning_statuses=commissioning_statuses)

@app.route('/mou/<int:mou_id>/edit', methods=['GET', 'POST'])
def edit_mou(mou_id):
    mou = MoU.query.get_or_404(mou_id)
    types = MasterValue.query.filter_by(category='type').order_by(MasterValue.name).all()
    sectors = MasterValue.query.filter_by(category='sector').order_by(MasterValue.name).all()
    countries = MasterValue.query.filter_by(category='country').order_by(MasterValue.name).all()
    states = MasterValue.query.filter_by(category='state').order_by(MasterValue.name).all()
    districts = MasterValue.query.filter_by(category='district').order_by(MasterValue.name).all()
    stages = MasterValue.query.filter_by(category='stage').order_by(MasterValue.name).all()
    land_statuses = MasterValue.query.filter_by(category='land_status').order_by(MasterValue.name).all()
    clearance_statuses = MasterValue.query.filter_by(category='clearance_status').order_by(MasterValue.name).all()
    commissioning_statuses = MasterValue.query.filter_by(category='commissioning_status').order_by(MasterValue.name).all()
    if request.method == 'POST':
        # Validate and update MoU
        mou.title = request.form.get('Title', '').strip()
        mou.company_name = request.form.get('company_name', '').strip()
        mou.type_id = request.form.get('type_id', type=int)
        mou.sector_id = request.form.get('sector_id', type=int)
        mou.investment = request.form.get('investment', type=float)
        mou.country_id = request.form.get('country_id', type=int)
        mou.state_id = request.form.get('state_id', type=int)
        mou.district_id = request.form.get('district_id', type=int)
        mou.stage_id = request.form.get('stage_id', type=int)
        mou.land_status_id = request.form.get('land_status_id', type=int)
        mou.clearance_status_id = request.form.get('clearance_status_id', type=int)
        mou.commissioning_status_id = request.form.get('commissioning_status_id', type=int)
        mou.start_date = datetime.strptime(request.form.get('start_date', ''), '%d-%m-%Y').date()
        mou.end_date = datetime.strptime(request.form.get('end_date', ''), '%d-%m-%Y').date()
        mou.remarks = request.form.get('remarks', '').strip()
        db.session.commit()
        # Audit log
        db.session.add(AuditLog(
            action_type='UPDATE',
            mou_id=mou.id,
            description=f"Updated MoU: {mou.title} for {mou.company_name}"
        ))
        db.session.commit()
        # Only generate documents if requested
        if request.form.get('generate_docs'):
            template_files = os.listdir(UPLOAD_FOLDER)
            template_file = next((f for f in template_files if f.endswith('.docx')), None)
            mapping_file = next((f for f in template_files if f.endswith('.mapping.json')), None)
            if template_file and mapping_file:
                template_path = os.path.join(UPLOAD_FOLDER, template_file)
                mapping_path = os.path.join(UPLOAD_FOLDER, mapping_file)
                with open(mapping_path, 'r') as f:
                    mapping = json.load(f)
                doc = Document(template_path)
                # Replace placeholders in paragraphs
                for para in doc.paragraphs:
                    for placeholder, field in mapping.items():
                        value = getattr(mou, field, '')
                        if isinstance(value, MasterValue):
                            value = value.name
                        if isinstance(value, datetime):
                            value = value.strftime('%d-%m-%Y')
                        para.text = para.text.replace(f'{{{{{placeholder}}}}}', str(value))
                # Replace placeholders in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for placeholder, field in mapping.items():
                                value = getattr(mou, field, '')
                                if isinstance(value, MasterValue):
                                    value = value.name
                                if isinstance(value, datetime):
                                    value = value.strftime('%d-%m-%Y')
                                cell.text = cell.text.replace(f'{{{{{placeholder}}}}}', str(value))
                word_filename = f'mou_{mou.id}.docx'
                word_path = os.path.join(generated_docs_folder, word_filename)
                if os.path.exists(word_path):
                    os.remove(word_path)
                doc.save(word_path)
                # Convert to PDF (cross-platform)
                pdf_filename = f'mou_{mou.id}.pdf'
                pdf_path = os.path.join(generated_docs_folder, pdf_filename)
                try:
                    if sys.platform.startswith('win'):
                        from docx2pdf import convert
                        convert(word_path, pdf_path)
                    else:
                        import pypandoc
                        output = pypandoc.convert_file(word_path, 'pdf', outputfile=pdf_path)
                        if output != "":
                            raise Exception(f"Pandoc error: {output}")
                    if not os.path.exists(pdf_path):
                        flash('PDF generation failed. Please check docx2pdf/pypandoc installation and permissions.', 'error')
                    else:
                        flash('MoU documents generated and replaced!', 'success')
                except Exception as e:
                    print('PDF conversion error:', e)
                    flash(f'PDF conversion error: {e}', 'error')
            else:
                flash('Template or mapping file not found.', 'error')
        else:
            flash('MoU updated!', 'success')
    return render_template('form.html', mou=mou, types=types, sectors=sectors, countries=countries, states=states, districts=districts, stages=stages, land_statuses=land_statuses, clearance_statuses=clearance_statuses, commissioning_statuses=commissioning_statuses)

@app.route('/download/mou/<int:mou_id>/<filetype>')
def download_mou_doc(mou_id, filetype):
    filename = f'mou_{mou_id}.{filetype}'
    filepath = os.path.abspath(os.path.join(generated_docs_folder, filename))
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    flash('File not found.', 'error')
    return redirect(url_for('list_mous'))

@app.route('/analytics')
def analytics():
    # Filters
    country_id = request.args.get('country', '').strip()
    sector_id = request.args.get('sector', '').strip()
    stage_id = request.args.get('stage', '').strip()
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()
    query = MoU.query
    if country_id:
        query = query.filter(MoU.country_id == int(country_id))
    if sector_id:
        query = query.filter(MoU.sector_id == int(sector_id))
    if stage_id:
        query = query.filter(MoU.stage_id == int(stage_id))
    if start_date:
        query = query.filter(MoU.start_date >= start_date)
    if end_date:
        query = query.filter(MoU.end_date <= end_date)
    mous = query.all()
    # Summary Cards
    total_mous = len(mous)
    total_investment = sum(mou.investment for mou in mous)
    active_mous = len([mou for mou in mous if mou.end_date >= datetime.now().date()])
    expiring_soon = len([mou for mou in mous if mou.end_date >= datetime.now().date() and mou.end_date <= (datetime.now().date().replace(day=1) + timedelta(days=30))])
    summary = {
        'total_mous': total_mous,
        'total_investment': total_investment,
        'active_mous': active_mous,
        'expiring_soon': expiring_soon
    }
    # Chart Data
    # Investment by Sector (sorted)
    sectors = MasterValue.query.filter_by(category='sector').order_by(MasterValue.name).all()
    sector_investments = [(s.name, sum(mou.investment for mou in mous if mou.sector and mou.sector.name == s.name)) for s in sectors]
    sector_investments.sort(key=lambda x: x[1], reverse=True)
    sector_labels = [x[0] for x in sector_investments]
    sector_data = [x[1] for x in sector_investments]
    investment_by_sector = {'labels': sector_labels, 'data': sector_data}
    # MoU Count by Country (sorted in JS)
    countries = MasterValue.query.filter_by(category='country').order_by(MasterValue.name).all()
    country_labels = [c.name for c in countries]
    country_data = [len([mou for mou in mous if mou.country and mou.country.name == c.name]) for c in countries]
    mou_count_by_country = {'labels': country_labels, 'data': country_data}
    # Stage Distribution (sorted in JS)
    stages = MasterValue.query.filter_by(category='stage').order_by(MasterValue.name).all()
    stage_labels = [s.name for s in stages]
    stage_data = [len([mou for mou in mous if mou.stage and mou.stage.name == s.name]) for s in stages]
    stage_distribution = {'labels': stage_labels, 'data': stage_data}
    # Start Date Trends (MoUs by month)
    from collections import Counter
    start_months = [mou.start_date.strftime('%Y-%m') for mou in mous if mou.start_date]
    start_counter = Counter(start_months)
    start_labels = sorted(start_counter.keys())
    start_data = [start_counter[month] for month in start_labels]
    start_date_trends = {'labels': start_labels, 'data': start_data}
    return render_template('analytics.html',
        countries=countries,
        sectors=sectors,
        stages=stages,
        investment_by_sector=investment_by_sector,
        mou_count_by_country=mou_count_by_country,
        stage_distribution=stage_distribution,
        start_date_trends=start_date_trends,
        summary=summary,
        request=request)

@app.route('/audit')
def audit():
    logs = AuditLog.query.order_by(AuditLog.timestamp.desc()).all()
    return render_template('audit.html', logs=logs)

@app.route('/settings', methods=['GET', 'POST'])
def settings():
    # Fetch all master values by category
    types = MasterValue.query.filter_by(category='type').order_by(MasterValue.name).all()
    sectors = MasterValue.query.filter_by(category='sector').order_by(MasterValue.name).all()
    countries = MasterValue.query.filter_by(category='country').order_by(MasterValue.name).all()
    states = MasterValue.query.filter_by(category='state').order_by(MasterValue.name).all()
    districts = MasterValue.query.filter_by(category='district').order_by(MasterValue.name).all()
    stages = MasterValue.query.filter_by(category='stage').order_by(MasterValue.name).all()
    land_statuses = MasterValue.query.filter_by(category='land_status').order_by(MasterValue.name).all()
    clearance_statuses = MasterValue.query.filter_by(category='clearance_status').order_by(MasterValue.name).all()
    commissioning_statuses = MasterValue.query.filter_by(category='commissioning_status').order_by(MasterValue.name).all()
    return render_template('settings.html', types=types, sectors=sectors, countries=countries, states=states, districts=districts, stages=stages, land_statuses=land_statuses, clearance_statuses=clearance_statuses, commissioning_statuses=commissioning_statuses)

@app.route('/import', methods=['GET', 'POST'])
def import_excel():
    if request.method == 'POST':
        import os
        from werkzeug.utils import secure_filename
        file = request.files.get('excel_file')
        if not file:
            flash('No file uploaded.', 'error')
            return redirect(url_for('import_excel'))
        filename = secure_filename(file.filename)
        filepath = os.path.join('uploads', filename)
        os.makedirs('uploads', exist_ok=True)
        file.save(filepath)
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active
        # Assuming first row is header
        for row in ws.iter_rows(min_row=2, values_only=True):
            (title, company_name, type_name, sector_name, investment, country_name, state_name, district_name, stage_name, land_status_name, clearance_status_name, commissioning_status_name, start_date, end_date, remarks) = row
            # Get or create master values
            type_obj = MasterValue.query.filter_by(name=type_name, category='type').first() or MasterValue(name=type_name, category='type')
            if not type_obj.id: db.session.add(type_obj)
            sector_obj = MasterValue.query.filter_by(name=sector_name, category='sector').first() or MasterValue(name=sector_name, category='sector')
            if not sector_obj.id: db.session.add(sector_obj)
            country_obj = MasterValue.query.filter_by(name=country_name, category='country').first() or MasterValue(name=country_name, category='country')
            if not country_obj.id: db.session.add(country_obj)
            state_obj = MasterValue.query.filter_by(name=state_name, category='state').first() or MasterValue(name=state_name, category='state')
            if not state_obj.id: db.session.add(state_obj)
            district_obj = MasterValue.query.filter_by(name=district_name, category='district').first() or MasterValue(name=district_name, category='district')
            if not district_obj.id: db.session.add(district_obj)
            stage_obj = MasterValue.query.filter_by(name=stage_name, category='stage').first() or MasterValue(name=stage_name, category='stage')
            if not stage_obj.id: db.session.add(stage_obj)
            land_status_obj = MasterValue.query.filter_by(name=land_status_name, category='land_status').first() or MasterValue(name=land_status_name, category='land_status')
            if not land_status_obj.id: db.session.add(land_status_obj)
            clearance_status_obj = MasterValue.query.filter_by(name=clearance_status_name, category='clearance_status').first() or MasterValue(name=clearance_status_name, category='clearance_status')
            if not clearance_status_obj.id: db.session.add(clearance_status_obj)
            commissioning_status_obj = MasterValue.query.filter_by(name=commissioning_status_name, category='commissioning_status').first() or MasterValue(name=commissioning_status_name, category='commissioning_status')
            if not commissioning_status_obj.id: db.session.add(commissioning_status_obj)
            db.session.flush() # Assign IDs
            mou = MoU(
                title=title,
                company_name=company_name,
                type_id=type_obj.id,
                sector_id=sector_obj.id,
                investment=investment,
                country_id=country_obj.id,
                state_id=state_obj.id,
                district_id=district_obj.id,
                stage_id=stage_obj.id,
                land_status_id=land_status_obj.id,
                clearance_status_id=clearance_status_obj.id,
                commissioning_status_id=commissioning_status_obj.id,
                start_date=start_date,
                end_date=end_date,
                remarks=remarks
            )
            db.session.add(mou)
            db.session.flush()
            db.session.add(AuditLog(
                action_type='IMPORT',
                mou_id=mou.id,
                description=f"Imported MoU: {title} for {company_name}"
            ))
        db.session.commit()
        flash('Excel import completed successfully!', 'success')
        return redirect(url_for('list_mous'))
    return render_template('import.html')

@app.route('/settings/<master>/add', methods=['POST'])
def add_master_value(master):
    name = request.form.get('name', '').strip()
    if not name:
        flash('Name is required.', 'error')
        return redirect(url_for('settings'))
    model = get_master_model(master)
    # Fix: Pass category when creating MasterValue
    if model.query.filter_by(name=name, category=master).first():
        flash(f'{master.title()} already exists.', 'error')
        return redirect(url_for('settings'))
    db.session.add(model(name=name, category=master))
    db.session.commit()
    flash(f'{master.title()} added.', 'success')
    return redirect(url_for('settings'))

@app.route('/settings/<master>/<int:id>/edit', methods=['POST'])
def edit_master_value(master, id):
    name = request.form.get('name', '').strip()
    model = get_master_model(master)
    item = model.query.get_or_404(id)
    item.name = name
    item.category = master  # Ensure category is set
    db.session.commit()
    flash(f'{master.title()} updated.', 'success')
    return redirect(url_for('settings'))

@app.route('/settings/<master>/<int:id>/delete', methods=['POST'])
def delete_master_value(master, id):
    model = get_master_model(master)
    item = model.query.get_or_404(id)
    db.session.delete(item)
    db.session.commit()
    flash(f'{master.title()} deleted.', 'success')
    return redirect(url_for('settings'))

@app.route('/template/upload', methods=['GET', 'POST'])
def upload_word_template():
    if request.method == 'POST':
        file = request.files.get('word_file')
        if not file or not file.filename.endswith('.docx'):
            flash('Please upload a valid .docx file.', 'error')
            return redirect(url_for('upload_word_template'))
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        # Remove any existing .docx files in the folder
        for f in os.listdir(UPLOAD_FOLDER):
            if f.endswith('.docx'):
                os.remove(os.path.join(UPLOAD_FOLDER, f))
        file.save(filepath)
        # Parse placeholders from Word file
        doc = Document(filepath)
        placeholders = set()
        import re
        # Check paragraphs
        for para in doc.paragraphs:
            found = re.findall(r'{{(.*?)}}', para.text)
            placeholders.update(found)
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found = re.findall(r'{{(.*?)}}', cell.text)
                    placeholders.update(found)
        placeholders = sorted(list(placeholders))
        # Show mapping UI
        return render_template('template_map.html', placeholders=placeholders, filename=filename)
    # List the single uploaded template for export
    template_files = [f for f in os.listdir(UPLOAD_FOLDER) if f.endswith('.docx')]
    return render_template('template_upload.html', template_files=template_files)

@app.route('/template/map/<filename>', methods=['POST'])
def save_template_mapping(filename):
    mapping = {}
    for key, value in request.form.items():
        if key.startswith('mapping_') and value:
            placeholder = key.replace('mapping_', '')
            mapping[placeholder] = value
    # Save mapping as JSON file
    mapping_path = os.path.join(UPLOAD_FOLDER, filename + '.mapping.json')
    with open(mapping_path, 'w') as f:
        json.dump(mapping, f)
    flash('Template mapping saved successfully!', 'success')
    return redirect(url_for('upload_word_template'))

@app.route('/template/export/<filename>')
def export_word_template(filename):
    filepath = os.path.abspath(os.path.join(UPLOAD_FOLDER, filename))
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    flash('Template file not found.', 'error')
    return redirect(url_for('upload_word_template'))

def get_master_model(master):
    return {
        'type': MasterValue,
        'sector': MasterValue,
        'country': MasterValue,
        'state': MasterValue,
        'district': MasterValue,
        'stage': MasterValue,
        'land_status': MasterValue,
        'clearance_status': MasterValue,
        'commissioning_status': MasterValue
    }[master]

@app.route('/import/template')
def download_excel_template():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([
        'Title', 'Company Name', 'Type', 'Sector', 'Investment', 'Country', 'State', 'District', 'Stage', 'Land Status', 'Clearance Status', 'Commissioning Status', 'Start Date', 'End Date', 'Remarks'
    ])
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, download_name='mou_template.xlsx', as_attachment=True)

@app.route('/mous/export')
def export_mous():
    # Filtering (same as list_mous)
    query = MoU.query
    company = request.args.get('company', '').strip()
    country_id = request.args.get('country', '').strip()
    type_id = request.args.get('type', '').strip()
    sector_id = request.args.get('sector', '').strip()
    if company:
        query = query.filter(MoU.company_name.ilike(f'%{company}%'))
    if country_id:
        query = query.filter(MoU.country_id == int(country_id))
    if type_id:
        query = query.filter(MoU.type_id == int(type_id))
    if sector_id:
        query = query.filter(MoU.sector_id == int(sector_id))
    # Sorting
    sort = request.args.get('sort', 'id')
    direction = request.args.get('direction', 'asc')
    if hasattr(MoU, sort):
        sort_col = getattr(MoU, sort)
        if direction == 'desc':
            sort_col = sort_col.desc()
        query = query.order_by(sort_col)
    # Pagination (only current page)
    page = request.args.get('page', 1, type=int)
    per_page = 10
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    mous = pagination.items
    # Export to Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([
        'ID', 'Title', 'Company Name', 'Type', 'Sector', 'Investment', 'Country', 'State', 'District', 'Stage', 'Start Date', 'End Date', 'Remarks'
    ])
    for mou in mous:
        ws.append([
            mou.id,
            mou.title,
            mou.company_name,
            mou.type.name if mou.type else '',
            mou.sector.name if mou.sector else '',
            mou.investment,
            mou.country.name if mou.country else '',
            mou.state.name if mou.state else '',
            mou.district.name if mou.district else '',
            mou.stage.name if mou.stage else '',
            mou.start_date.strftime('%d-%m-%Y') if mou.start_date else '',
            mou.end_date.strftime('%d-%m-%Y') if mou.end_date else '',
            mou.remarks or ''
        ])
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, download_name='mou_listing.xlsx', as_attachment=True)

@app.route('/mou/<int:mou_id>/delete', methods=['POST'])
def delete_mou(mou_id):
    mou = MoU.query.get_or_404(mou_id)
    db.session.delete(mou)
    db.session.add(AuditLog(
        action_type='DELETE',
        mou_id=mou_id,
        description=f"Deleted MoU ID: {mou_id}"
    ))
    db.session.commit()
    flash('MoU deleted successfully!', 'success')
    return redirect(url_for('list_mous'))

@app.route('/db-import-export', methods=['GET', 'POST'])
def db_import_export():
    instance_db_path = os.path.abspath(os.path.join('instance', 'moutracker.db'))
    message = None
    if request.method == 'POST':
        if 'export_db' in request.form:
            if os.path.exists(instance_db_path):
                return send_file(instance_db_path, as_attachment=True, download_name='moutracker.db')
            else:
                message = 'Database file not found in instance folder. Please create some data first.'
        elif 'import_db' in request.form:
            file = request.files.get('db_file')
            if file and file.filename.endswith('.db'):
                file.save(instance_db_path)
                message = 'Database imported successfully to instance folder.'
            else:
                message = 'Please upload a valid .db file.'
    return render_template('db_import_export.html', db_path=instance_db_path, message=message)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()  # Create database and tables if they do not exist
    app.run(debug=True)
